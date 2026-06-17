from __future__ import annotations

import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "REPORT.docx"


def clean_docx_metadata(path: Path) -> None:
    core_properties = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>Отчет по проекту</dc:title><dc:subject>Банковские транзакции</dc:subject><dc:creator/><cp:keywords/><dc:description/><cp:lastModifiedBy/><cp:revision>1</cp:revision><dcterms:created xsi:type="dcterms:W3CDTF">2026-06-11T00:00:00Z</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">2026-06-11T00:00:00Z</dcterms:modified><cp:category/></cp:coreProperties>"""

    handle, temp_name = tempfile.mkstemp(suffix=".docx")
    Path(temp_name).unlink()
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_name, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = core_properties.encode("utf-8") if item.filename == "docProps/core.xml" else source.read(item.filename)
                target.writestr(item, data)
        shutil.move(temp_name, path)
    finally:
        try:
            Path(temp_name).unlink()
        except FileNotFoundError:
            pass
        try:
            import os

            os.close(handle)
        except OSError:
            pass


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_margins(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build() -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Отчет по проекту\nБанковские транзакции")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Вариант 9. C# / .NET 8 / EF Core / SQLite / Avalonia / xUnit")

    doc.add_heading("1. Предметная область", level=1)
    doc.add_paragraph(
        "Система регистрирует банковские транзакции и обновляет счета клиентов. "
        "Поддерживаются зарплатные, валютные и накопительные счета. Клиенты делятся "
        "на физических и юридических лиц. Банк работает с банками-партнерами, включая "
        "зарубежные, а комиссия рассчитывается по проценту банка-партнера."
    )

    doc.add_heading("2. Состав решения", level=1)
    add_table(
        doc,
        ["Компонент", "Назначение"],
        [
            ["BankTransactions.Server", "Консольный HTTP API, авторизация, бизнес-логика, SQLite"],
            ["BankTransactions.Client", "Оконный Avalonia-клиент оператора банка"],
            ["BankTransactions.Shared", "DTO-контракты клиента, сервера и тестов"],
            ["BankTransactions.Tests", "Unit и сквозные тесты xUnit"],
        ],
    )

    doc.add_heading("3. Архитектура", level=1)
    doc.add_paragraph(
        "Оконный клиент не обращается к базе данных напрямую. Все действия оператора "
        "передаются на сервер по HTTP. Сервер проверяет bearer-токен, валидирует данные, "
        "выбирает режим доступа к БД и возвращает DTO."
    )
    add_bullets(
        doc,
        [
            "GUI -> HTTP JSON -> ASP.NET Core API -> BankingService -> EF Core ORM -> SQLite.",
            "GUI -> HTTP JSON -> ASP.NET Core API -> BankingService -> SQL commands -> SQLite.",
            "Режим выбирается параметром mode=orm или mode=sql.",
        ],
    )

    doc.add_heading("4. База данных", level=1)
    add_table(
        doc,
        ["Таблица", "Ключевые поля", "Назначение"],
        [
            ["Users", "Username, PasswordHash, Role", "Пользователи и авторизация"],
            ["Customers", "CustomerType, FullName, TaxId", "Физические и юридические лица"],
            ["PartnerBanks", "SwiftCode, IsForeign, FeePercent", "Банки-партнеры и комиссия"],
            ["Accounts", "AccountNumber, AccountType, Currency, Balance", "Счета клиентов"],
            ["Transactions", "FromAccountId, ToAccountId, Amount, FeeAmount", "История переводов"],
        ],
    )

    doc.add_heading("5. Реализованные функции", level=1)
    add_table(
        doc,
        ["Функция", "Реализация"],
        [
            ["Авторизация", "Логин admin/admin123, bearer-токен, middleware-защита API"],
            ["Хеш паролей", "PBKDF2 + соль + SHA-256, открытый пароль не хранится"],
            ["CRUD клиентов", "Добавление, изменение, удаление, поиск через ORM и SQL"],
            ["CRUD счетов", "Добавление, изменение, удаление, поиск через ORM и SQL"],
            ["CRUD партнеров", "Добавление, изменение, удаление, поиск через ORM и SQL"],
            ["Транзакции", "Создание перевода, комиссия, обновление балансов, поиск"],
        ],
    )

    doc.add_heading("6. Фрагменты кода", level=1)
    doc.add_paragraph("Пример SQL-поиска клиентов:")
    add_code(
        doc,
        "SELECT Id, CustomerType, FullName, TaxId, Email, Phone, CreatedAtUtc\n"
        "FROM Customers\n"
        "WHERE $pattern = '%%' OR FullName LIKE $pattern OR TaxId LIKE $pattern\n"
        "ORDER BY FullName"
    )
    doc.add_paragraph("Пример ORM-поиска клиентов:")
    add_code(
        doc,
        "query = query.Where(customer =>\n"
        "    customer.FullName.Contains(value)\n"
        "    || customer.TaxId.Contains(value)\n"
        "    || customer.Email.Contains(value));"
    )

    doc.add_heading("7. Диаграммы", level=1)
    add_bullets(
        doc,
        [
            "IDEF0: docs/diagrams/context-idef0.mmd",
            "IDEF3: docs/diagrams/idef3-process.mmd",
            "DFD: docs/diagrams/dfd.mmd",
            "UML Use Case: docs/diagrams/use-cases.mmd",
            "UML Class Diagram: docs/diagrams/class-diagram.mmd",
            "UML Sequence Diagram: docs/diagrams/sequence-transfer.mmd",
            "ER-схема БД: docs/diagrams/db-schema.mmd",
        ],
    )

    doc.add_heading("8. Тестирование", level=1)
    add_table(
        doc,
        ["Тест", "Тип", "Проверка", "Результат"],
        [
            ["PasswordHasherTests", "Unit", "Хеш не хранит пароль, verify работает", "Passed"],
            ["DataAccessModeParser", "Unit", "Выбор режима SQL/ORM, 5 случаев", "Passed"],
            ["AccountNumberGenerator", "Unit", "Префикс, длина и цифровой формат счета", "Passed"],
            ["CustomerCrud", "Сквозной", "SQL create, ORM update, search, delete, SQLite", "Passed"],
            ["EndToEndTransfer", "Сквозной", "Создание счетов, перевод, комиссия, баланс", "Passed"],
        ],
    )

    doc.add_heading("9. Запуск", level=1)
    add_code(doc, "dotnet build BankTransactions.sln\ndotnet test BankTransactions.sln")
    add_code(
        doc,
        "dotnet run --project src/BankTransactions.Server --urls http://localhost:5055\n"
        "dotnet run --project src/BankTransactions.Client"
    )

    doc.add_heading("10. Вывод", level=1)
    doc.add_paragraph(
        "Проект реализует требования задания: серверная консольная часть с БД, оконный "
        "клиент, HTTP-взаимодействие, авторизация, хеширование паролей, CRUD/search "
        "через SQL и ORM, тесты, диаграммы проектирования и схема базы данных."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    clean_docx_metadata(OUT)


if __name__ == "__main__":
    build()
